import io

from docx import Document
from src.infrastructure.parsing.text_extraction import DocumentTextExtractor


def test_extract_text_from_docx_round_trips_paragraphs() -> None:
    document = Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("Senior Engineer")
    buffer = io.BytesIO()
    document.save(buffer)

    text = DocumentTextExtractor().extract_text(buffer.getvalue(), "docx")

    assert "Jane Doe" in text
    assert "Senior Engineer" in text
